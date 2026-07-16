"""Bounded in-memory parsing for PDF, DOCX, and UTF-8 TXT uploads."""

from __future__ import annotations

import hashlib
import io
import re
import unicodedata
import zipfile
from pathlib import PurePosixPath
from typing import Literal, cast

from docx import Document as DocxDocument
from pdfminer.high_level import extract_text as pdfminer_extract
from pdfminer.pdfdocument import PDFDocument, PDFPasswordIncorrect
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser

from skillmap.config.settings import Settings, get_settings
from skillmap.core.exceptions import UserFacingError
from skillmap.domain.models import ParsedDocument

ALLOWED_TYPES = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt": "text/plain",
}
_SAFE_NAME_RE = re.compile(r"[^A-Za-z0-9._ -]+")
_CONTROL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]")
_WHITESPACE_RE = re.compile(r"[ \t]+")


class DocumentValidationError(UserFacingError):
    def __init__(self, category: str) -> None:
        super().__init__(
            "We could not process this document. Confirm that it is a valid PDF, DOCX or TXT file under the configured size limit.",
            category=category,
        )


def sanitize_filename(filename: str) -> str:
    name = filename.replace("\\", "/").rsplit("/", 1)[-1]
    name = _SAFE_NAME_RE.sub("_", unicodedata.normalize("NFKC", name)).strip(" .")
    return (name or "upload")[:120]


def _extension(filename: str) -> str:
    dot = filename.rfind(".")
    return filename[dot:].lower() if dot >= 0 else ""


def _validate_type(data: bytes | bytearray, filename: str, content_type: str) -> str:
    extension = _extension(filename)
    expected_mime = ALLOWED_TYPES.get(extension)
    if expected_mime is None:
        raise DocumentValidationError("unsupported_extension")
    actual_mime = content_type.partition(";")[0].strip().lower()
    if actual_mime != expected_mime:
        raise DocumentValidationError("mime_mismatch")
    if extension == ".pdf" and not data.startswith(b"%PDF-"):
        raise DocumentValidationError("signature_mismatch")
    if extension == ".docx" and not data.startswith(b"PK\x03\x04"):
        raise DocumentValidationError("signature_mismatch")
    if extension == ".txt" and b"\x00" in data[:4096]:
        raise DocumentValidationError("invalid_text_file")
    return extension[1:]


def _safe_docx_archive(data: bytes | bytearray) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            infos = archive.infolist()
            if len(infos) > 250:
                raise DocumentValidationError("archive_file_limit")
            total_size = 0
            for info in infos:
                normalized = info.filename.replace("\\", "/")
                path = PurePosixPath(normalized)
                if path.is_absolute() or ".." in path.parts:
                    raise DocumentValidationError("archive_path_traversal")
                if (info.external_attr >> 16) & 0o170000 == 0o120000:
                    raise DocumentValidationError("archive_symlink")
                total_size += info.file_size
                if info.file_size > 0 and info.file_size / max(info.compress_size, 1) > 100:
                    raise DocumentValidationError("archive_expansion_ratio")
            if total_size > 20 * 1024 * 1024:
                raise DocumentValidationError("archive_expansion_limit")
            content_types = archive.read("[Content_Types].xml").lower()
            if b"macroenabled" in content_types or b"vbaproject" in content_types:
                raise DocumentValidationError("macro_enabled_document")
    except KeyError as exc:
        raise DocumentValidationError("invalid_docx") from exc
    except zipfile.BadZipFile as exc:
        raise DocumentValidationError("invalid_docx") from exc


def _extract_pdf(data: bytes | bytearray, max_pages: int) -> tuple[str, int]:
    stream = io.BytesIO(data)
    try:
        document = PDFDocument(PDFParser(stream))
        if not document.is_extractable:
            raise DocumentValidationError("encrypted_pdf")
        page_count = sum(1 for _ in PDFPage.create_pages(document))
    except PDFPasswordIncorrect as exc:
        raise DocumentValidationError("encrypted_pdf") from exc
    except DocumentValidationError:
        raise
    except Exception as exc:
        raise DocumentValidationError("invalid_pdf") from exc
    if page_count == 0 or page_count > max_pages:
        raise DocumentValidationError("pdf_page_limit")
    try:
        return pdfminer_extract(io.BytesIO(data), maxpages=max_pages), page_count
    except Exception as exc:
        raise DocumentValidationError("pdf_parse_error") from exc


def _extract_docx(data: bytes | bytearray) -> str:
    _safe_docx_archive(data)
    try:
        document = DocxDocument(io.BytesIO(data))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                paragraphs.append(" ".join(cell.text for cell in row.cells))
        return "\n".join(paragraphs)
    except DocumentValidationError:
        raise
    except Exception as exc:
        raise DocumentValidationError("docx_parse_error") from exc


def _extract_txt(data: bytes | bytearray) -> str:
    try:
        return bytes(data).decode("utf-8-sig", errors="strict")
    except UnicodeDecodeError as exc:
        raise DocumentValidationError("invalid_text_encoding") from exc


def _normalize_text(text: str, max_chars: int) -> str:
    text = unicodedata.normalize("NFKC", text)
    text = _CONTROL_RE.sub(" ", text)
    text = "\n".join(_WHITESPACE_RE.sub(" ", line).strip() for line in text.splitlines())
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        raise DocumentValidationError("empty_document")
    if len(text) > max_chars:
        raise DocumentValidationError("extracted_text_limit")
    return text


def parse_document(
    data: bytes | bytearray,
    filename: str,
    content_type: str,
    settings: Settings | None = None,
) -> ParsedDocument:
    settings = settings or get_settings()
    if not data:
        raise DocumentValidationError("empty_file")
    if len(data) > settings.max_resume_bytes:
        raise DocumentValidationError("file_size_limit")
    safe_name = sanitize_filename(filename)
    file_type = cast(
        Literal["pdf", "docx", "txt"],
        _validate_type(data, safe_name, content_type),
    )
    page_count: int | None = None
    if file_type == "pdf":
        text, page_count = _extract_pdf(data, settings.max_pdf_pages)
    elif file_type == "docx":
        text = _extract_docx(data)
    else:
        text = _extract_txt(data)
    return ParsedDocument(
        filename=safe_name,
        file_type=file_type,
        text=_normalize_text(text, settings.max_extracted_text_chars),
        size_bytes=len(data),
        sha256=hashlib.sha256(data).hexdigest(),
        page_count=page_count,
    )
